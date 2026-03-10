"""Build a Word document from OPENAPI_MCP_SPEC.md with rendered Mermaid figures."""

import re
import subprocess
import tempfile
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

SPEC = Path(__file__).parent / "OPENAPI_MCP_SPEC.md"
OUT = Path(__file__).parent / "OPENAPI_MCP_SPEC.docx"


def render_mermaid(code: str, index: int, tmp_dir: Path) -> Path:
    """Render a Mermaid code block to PNG via mmdc."""
    mmd_file = tmp_dir / f"diagram_{index}.mmd"
    png_file = tmp_dir / f"diagram_{index}.png"
    mmd_file.write_text(code, encoding="utf-8")
    result = subprocess.run(
        f'mmdc -i "{mmd_file}" -o "{png_file}" -b transparent -s 2',
        capture_output=True, text=True, shell=True,
    )
    if result.returncode != 0:
        print(f"  mmdc stderr: {result.stderr}")
        raise RuntimeError(result.stderr)
    if not png_file.exists():
        raise FileNotFoundError(f"mmdc did not produce {png_file}")
    return png_file


def parse_table(lines: list[str]) -> list[list[str]]:
    """Parse markdown table lines into rows of cells."""
    rows = []
    for line in lines:
        line = line.strip()
        if not line.startswith("|"):
            continue
        cells = [c.strip() for c in line.split("|")[1:-1]]
        # Skip separator rows (e.g., |---|---|)
        if all(re.match(r"^[-:]+$", c) for c in cells):
            continue
        rows.append(cells)
    return rows


def set_cell_shading(cell, color_hex: str):
    """Apply background shading to a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(
        qn("w:shd"),
        {qn("w:fill"): color_hex, qn("w:val"): "clear"},
    )
    shading.append(shading_elem)


def add_table(doc, rows: list[list[str]]):
    """Add a formatted table to the document."""
    if not rows:
        return
    n_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(cell_text)
            run.font.size = Pt(9)
            if i == 0:
                run.bold = True
                set_cell_shading(cell, "D9E2F3")


def add_code_block(doc, code: str):
    """Add a code block as a formatted paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(30, 30, 30)


def add_inline_formatting(paragraph, text: str):
    """Add text with inline bold (**) and code (`) formatting."""
    # Split on bold and code markers
    parts = re.split(r"(\*\*.*?\*\*|`[^`]+`)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(60, 60, 60)
        else:
            paragraph.add_run(part)


def build_docx():
    md = SPEC.read_text(encoding="utf-8")
    lines = md.split("\n")

    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    tmp_dir = Path(tempfile.mkdtemp())
    mermaid_idx = 0

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip horizontal rules
        if stripped == "---":
            i += 1
            continue

        # Mermaid code block
        if stripped == "```mermaid":
            mermaid_lines = []
            i += 1
            while i < len(lines) and lines[i].strip() != "```":
                mermaid_lines.append(lines[i])
                i += 1
            i += 1  # skip closing ```
            code = "\n".join(mermaid_lines)
            try:
                png_path = render_mermaid(code, mermaid_idx, tmp_dir)
                mermaid_idx += 1
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(str(png_path), width=Inches(5.5))
            except Exception as e:
                print(f"  ERROR rendering diagram {mermaid_idx}: {e}")
                add_code_block(doc, f"[Mermaid render failed: {e}]\n{code}")
            continue

        # Generic code block
        if stripped.startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and lines[i].strip() != "```":
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing ```
            add_code_block(doc, "\n".join(code_lines))
            continue

        # Table
        if stripped.startswith("|") and i + 1 < len(lines) and lines[i + 1].strip().startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            rows = parse_table(table_lines)
            add_table(doc, rows)
            continue

        # Headings
        if stripped.startswith("# ") and not stripped.startswith("## "):
            p = doc.add_heading(stripped[2:], level=0)
            i += 1
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=1)
            i += 1
            continue
        if stripped.startswith("#### "):
            doc.add_heading(stripped[5:], level=3)
            i += 1
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=2)
            i += 1
            continue

        # Blockquote
        if stripped.startswith("> "):
            p = doc.add_paragraph(style="Intense Quote")
            add_inline_formatting(p, stripped[2:])
            i += 1
            continue

        # Bullet list
        if stripped.startswith("- [ ] ") or stripped.startswith("- [x] "):
            checked = stripped.startswith("- [x] ")
            prefix = "☑ " if checked else "☐ "
            p = doc.add_paragraph(style="List Bullet")
            add_inline_formatting(p, prefix + stripped[6:])
            i += 1
            continue
        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_formatting(p, stripped[2:])
            i += 1
            continue
        if re.match(r"^\d+\.\s", stripped):
            text = re.sub(r"^\d+\.\s+", "", stripped)
            p = doc.add_paragraph(style="List Number")
            add_inline_formatting(p, text)
            i += 1
            continue

        # Blank line
        if not stripped:
            i += 1
            continue

        # Normal paragraph
        p = doc.add_paragraph()
        add_inline_formatting(p, stripped)
        i += 1

    doc.save(str(OUT))
    print(f"Created: {OUT}")
    print(f"Rendered {mermaid_idx} Mermaid diagrams")


if __name__ == "__main__":
    build_docx()
